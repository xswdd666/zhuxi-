"""Endpoint-level checks for the Stage B project and review foundation."""

from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from fastapi.testclient import TestClient

from app import db
from app.main import app
from app.services import parser, rag, repositories

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover - python-docx is a normal dependency
    DOCX_AVAILABLE = False

try:
    import win32com.client as win32
    WORD_AVAILABLE = True
except Exception:  # pragma: no cover - pywin32 may not be installed
    WORD_AVAILABLE = False


class StageBApiTests(unittest.TestCase):
    def test_project_document_source_and_task_log_endpoints(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            with patch.object(db, "DATABASE_PATH", root / "test.sqlite3"), patch.object(db, "DATA_DIR", root / "data"), patch.object(db, "UPLOADS_DIR", root / "uploads"), patch.object(db, "OUTPUT_DIR", root / "output"), patch.object(db, "STATIC_DIR", root / "static"), patch.object(parser, "UPLOADS_DIR", root / "uploads"), patch.object(rag, "index_project_chunks", return_value={"rag_status": "ready", "indexed_count": 1}), patch.object(rag, "delete_document_vectors", return_value={"rag_status": "ready"}):
                with TestClient(app) as client:
                    created = client.post("/api/projects", json={"name": "接口验收项目"})
                    self.assertEqual(created.status_code, 201)
                    project_id = created.json()["id"]
                    updated = client.patch(f"/api/projects/{project_id}", json={"location": "杭州", "stage": "前期资料分析", "objective": "验证接口", "tags": ["滨水"]})
                    self.assertEqual(updated.status_code, 200)
                    self.assertEqual(updated.json()["tags"], ["滨水"])

                    upload = client.post(f"/api/projects/{project_id}/documents", files=[("files", ("brief.txt", "场地入口需要保留。", "text/plain"))])
                    self.assertEqual(upload.status_code, 200)
                    document = upload.json()["documents"][0]
                    self.assertNotIn("path", document)
                    document_id = document["id"]
                    chunks = client.get(f"/api/documents/{document_id}/chunks").json()["chunks"]
                    self.assertEqual(len(chunks), 1)
                    source = client.get(f"/api/source-chunks/{chunks[0]['id']}")
                    self.assertEqual(source.status_code, 200)
                    self.assertNotIn("path", source.json()["source_chunk"])
                    self.assertIn("场地入口", client.get(f"/api/documents/{document_id}/preview").text)
                    self.assertEqual(client.post(f"/api/documents/{document_id}/retry").status_code, 200)

                    repositories.create_task_log(project_id, "test.stage_b", "completed", "接口检查完成")
                    self.assertEqual(client.get(f"/api/projects/{project_id}/task-logs").json()["logs"][0]["task_type"], "test.stage_b")
                    repositories.create_insight(project_id, "site_fact", "被引用的资料", "该资料正在被洞察引用。", [{"document_id": document_id, "file_name": "brief.txt", "locator": "文本片段 1", "quote": "场地入口", "source_chunk_id": chunks[0]["id"]}], 0.8)
                    blocked = client.delete(f"/api/documents/{document_id}")
                    self.assertEqual(blocked.status_code, 409)
                    self.assertEqual(blocked.json()["error"]["code"], "DOCUMENT_IN_USE")
                    with db.connection_scope() as connection:
                        connection.execute("DELETE FROM insight_cards WHERE project_id = ?", (project_id,))
                    self.assertTrue(client.delete(f"/api/documents/{document_id}").json()["deleted"])

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
    def test_docx_upload_is_parsed_and_chunked(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            buffer = io.BytesIO()
            doc = Document()
            doc.add_heading("示例 DOCX 任务书", level=1)
            doc.add_paragraph("项目：社区活动中心改造")
            doc.add_paragraph("约束：现状北侧保留两层建筑")
            doc.save(buffer)
            buffer.seek(0)

            with patch.object(db, "DATABASE_PATH", root / "test.sqlite3"), patch.object(db, "DATA_DIR", root / "data"), patch.object(db, "UPLOADS_DIR", root / "uploads"), patch.object(db, "OUTPUT_DIR", root / "output"), patch.object(db, "STATIC_DIR", root / "static"), patch.object(parser, "UPLOADS_DIR", root / "uploads"), patch.object(rag, "index_project_chunks", return_value={"rag_status": "degraded", "indexed_count": 0, "reason": "test"}):
                with TestClient(app) as client:
                    project_id = client.post("/api/projects", json={"name": "DOCX 验收项目"}).json()["id"]
                    upload = client.post(f"/api/projects/{project_id}/documents", files=[("files", ("brief.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))])
                    self.assertEqual(upload.status_code, 200)
                    document = upload.json()["documents"][0]
                    self.assertEqual(document["file_type"], "docx")
                    self.assertTrue(document["parse_status"].startswith("parsed"))
                    chunks = client.get(f"/api/documents/{document['id']}/chunks").json()["chunks"]
                    self.assertTrue(any("社区活动中心" in chunk["content"] for chunk in chunks))

    @unittest.skipUnless(WORD_AVAILABLE, "Word COM not available")
    def test_doc_upload_is_parsed_via_word_com(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            doc_path = root / "brief.doc"
            try:
                word = win32.Dispatch("Word.Application")
            except Exception as exc:  # Word can be installed but unavailable to a non-interactive session.
                self.skipTest(f"Word COM session unavailable: {type(exc).__name__}")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Add()
            doc.Paragraphs[0].Range.Text = "项目：旧版 DOC 测试\n"
            doc.Paragraphs.Add().Range.Text = "约束：高度不超过 12 米"
            doc.SaveAs(str(doc_path.resolve()), FileFormat=0)  # 0 = wdFormatDocument
            doc.Close(SaveChanges=False)
            word.Quit(SaveChanges=-1)

            status, chunks = parser.parse_file(doc_path, "doc")
            self.assertEqual(status, "parsed")
            self.assertTrue(any("旧版 DOC 测试" in chunk[1] for chunk in chunks))


if __name__ == "__main__":
    unittest.main()
